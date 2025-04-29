"""Main Streamlit application for Book to Essay Generator."""
import streamlit as st
from pathlib import Path
import tempfile
import os
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from src.book_to_essay.ai_book_to_essay_generator import AIBookEssayGenerator
from src.book_to_essay.config import SUPPORTED_FORMATS, MAX_UPLOAD_SIZE_MB

st.set_page_config(
    page_title="Book to Essay Generator",
    page_icon="📚",
    layout="wide"
)

def create_docx(text):
    """Create a Word document from text."""
    doc = Document()
    doc.add_heading('Generated Essay', 0)
    doc.add_paragraph(text)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_pdf(text):
    """Create a PDF document from text."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []
    
    # Add title
    title = Paragraph("Generated Essay", styles['Title'])
    flowables.append(title)
    
    # Add content
    content = Paragraph(text, styles['Normal'])
    flowables.append(content)
    
    # Build PDF
    doc.build(flowables)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    st.title("📚 Book to Essay Generator")
    
    # Initialize session state
    if 'generator' not in st.session_state:
        st.session_state.generator = AIBookEssayGenerator()
    if 'uploaded_files' not in st.session_state:
        st.session_state.uploaded_files = []

    # Sidebar for file uploads and settings
    with st.sidebar:
        st.header("📥 Upload Books")
        uploaded_files = st.file_uploader(
            "Upload your books",
            accept_multiple_files=True,
            type=[fmt.replace('.', '') for fmt in SUPPORTED_FORMATS]
        )
        
        st.caption(f"Supported formats: {', '.join(SUPPORTED_FORMATS)}")
        st.caption(f"Max file size: {MAX_UPLOAD_SIZE_MB}MB")

    # Main content area
    col1, col2 = st.columns([2, 3])
    
    with col1:
        st.header("Essay Parameters")
        topic = st.text_area("Essay Topic/Prompt", height=100)
        word_limit = st.number_input("Word Limit", min_value=100, max_value=5000, value=500, step=100)
        style = st.selectbox("Writing Style", ["Academic", "Analytical", "Argumentative", "Expository"])
        
        if st.button("Generate Essay", type="primary"):
            if not uploaded_files:
                st.error("Please upload at least one book first!")
                return
            if not topic:
                st.error("Please provide an essay topic!")
                return
                
            with st.spinner("Processing books and generating essay..."):
                try:
                    # Create a temporary directory for uploaded files
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Process uploaded files
                        for uploaded_file in uploaded_files:
                            try:
                                # Create a temporary file path
                                temp_file_path = os.path.join(temp_dir, uploaded_file.name)
                                
                                # Save the uploaded file
                                with open(temp_file_path, 'wb') as f:
                                    f.write(uploaded_file.getbuffer())
                                
                                # Process the file based on its type
                                if temp_file_path.endswith('.pdf'):
                                    st.session_state.generator.load_pdf_file(temp_file_path)
                                elif temp_file_path.endswith('.txt'):
                                    st.session_state.generator.load_txt_file(temp_file_path)
                                elif temp_file_path.endswith('.epub'):
                                    st.session_state.generator.load_epub_file(temp_file_path)
                                    
                                st.success(f"Successfully processed {uploaded_file.name}")
                            except Exception as e:
                                st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                                continue
                        
                        # Generate essay
                        essay = st.session_state.generator.generate_essay(topic, word_limit)
                        st.session_state.essay = essay
                        st.session_state.quotes = st.session_state.generator.extract_quotes()
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
                    return

    with col2:
        st.header("Generated Essay")
        if 'essay' in st.session_state:
            st.markdown(st.session_state.essay)
            
            st.subheader("Supporting Quotes")
            for quote in st.session_state.quotes:
                st.markdown(f"> {quote}")
            
            # Logging for debugging
            import logging
            logging.info(f"Essay for download (first 100 chars): {st.session_state.essay[:100]}")
            logging.info(f"Quotes for download: {st.session_state.quotes}")
            
            # Always regenerate downloadable files from the latest essay
            col_word, col_pdf = st.columns(2)
            with col_word:
                logging.info("Generating downloadable files...")
                docx_data = create_docx(st.session_state.essay)
                logging.info(f"DOCX data generated for download, length: {len(docx_data)} bytes")
                st.download_button(
                    "Download as Word",
                    docx_data,
                    file_name="generated_essay.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                logging.info("Generated downloadable files.")
            with col_pdf:
                logging.info("Generating downloadable files...")
                pdf_data = create_pdf(st.session_state.essay)
                logging.info(f"PDF data generated for download, length: {len(pdf_data)} bytes")
                st.download_button(
                    "Download as PDF",
                    pdf_data,
                    file_name="generated_essay.pdf",
                    mime="application/pdf"
                )
                logging.info("Generated downloadable files.")

if __name__ == "__main__":
    main()
